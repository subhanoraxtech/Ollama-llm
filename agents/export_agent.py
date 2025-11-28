"""
Export Agent - handles file export and report generation.
"""
import pandas as pd
from typing import Dict, Any, Optional
import os
from .base_agent import BaseAgent
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from docx import Document
from docx.shared import Inches


class ExportAgent(BaseAgent):
    """Agent for exporting data and generating reports."""
    
    def __init__(self):
        super().__init__("ExportAgent")
    
    def get_available_operations(self) -> Dict[str, str]:
        """Get available export operations."""
        return {
            'export_to_csv': 'Export DataFrame to CSV file',
            'export_to_excel': 'Export DataFrame to Excel file',
            'export_to_json': 'Export DataFrame to JSON file',
            'generate_pdf_report': 'Generate PDF summary report',
            'generate_word_report': 'Generate Word document report'
        }
    
    def export_to_csv(
        self,
        df: pd.DataFrame,
        file_path: str,
        include_index: bool = False
    ) -> Dict[str, Any]:
        """
        Export DataFrame to CSV.
        
        Args:
            df: DataFrame to export
            file_path: Output file path
            include_index: Include index in export
            
        Returns:
            Result dictionary
        """
        try:
            is_valid, error = self.validate_dataframe(df)
            if not is_valid:
                return self.create_result(False, message=error)
            
            # Ensure directory exists
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            
            df.to_csv(file_path, index=include_index)
            
            file_size = os.path.getsize(file_path)
            
            self.log_operation('export_to_csv', {
                'file_path': file_path,
                'rows': len(df),
                'columns': len(df.columns),
                'file_size_kb': file_size / 1024
            })
            
            return self.create_result(
                success=True,
                message=f"Exported {len(df)} rows to {file_path}",
                metadata={
                    'file_path': file_path,
                    'rows': len(df),
                    'columns': len(df.columns),
                    'file_size_kb': round(file_size / 1024, 2)
                }
            )
        
        except Exception as e:
            return self.handle_error('export_to_csv', e)
    
    def export_to_excel(
        self,
        df: pd.DataFrame,
        file_path: str,
        sheet_name: str = 'Sheet1',
        include_index: bool = False
    ) -> Dict[str, Any]:
        """
        Export DataFrame to Excel.
        
        Args:
            df: DataFrame to export
            file_path: Output file path
            sheet_name: Name of the Excel sheet
            include_index: Include index in export
            
        Returns:
            Result dictionary
        """
        try:
            is_valid, error = self.validate_dataframe(df)
            if not is_valid:
                return self.create_result(False, message=error)
            
            # Ensure directory exists
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            
            # Create Excel writer
            with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
                df.to_excel(writer, sheet_name=sheet_name, index=include_index)
                
                # Auto-adjust column widths
                worksheet = writer.sheets[sheet_name]
                for column in worksheet.columns:
                    max_length = 0
                    column_letter = column[0].column_letter
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                    adjusted_width = min(max_length + 2, 50)
                    worksheet.column_dimensions[column_letter].width = adjusted_width
            
            file_size = os.path.getsize(file_path)
            
            self.log_operation('export_to_excel', {
                'file_path': file_path,
                'rows': len(df),
                'columns': len(df.columns),
                'file_size_kb': file_size / 1024
            })
            
            return self.create_result(
                success=True,
                message=f"Exported {len(df)} rows to {file_path}",
                metadata={
                    'file_path': file_path,
                    'rows': len(df),
                    'columns': len(df.columns),
                    'file_size_kb': round(file_size / 1024, 2)
                }
            )
        
        except Exception as e:
            return self.handle_error('export_to_excel', e)
    
    def export_to_json(
        self,
        df: pd.DataFrame,
        file_path: str,
        orient: str = 'records'
    ) -> Dict[str, Any]:
        """
        Export DataFrame to JSON.
        
        Args:
            df: DataFrame to export
            file_path: Output file path
            orient: JSON orientation ('records', 'index', 'columns', 'values')
            
        Returns:
            Result dictionary
        """
        try:
            is_valid, error = self.validate_dataframe(df)
            if not is_valid:
                return self.create_result(False, message=error)
            
            # Ensure directory exists
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            
            df.to_json(file_path, orient=orient, indent=2)
            
            file_size = os.path.getsize(file_path)
            
            self.log_operation('export_to_json', {
                'file_path': file_path,
                'rows': len(df),
                'columns': len(df.columns),
                'file_size_kb': file_size / 1024
            })
            
            return self.create_result(
                success=True,
                message=f"Exported {len(df)} rows to {file_path}",
                metadata={
                    'file_path': file_path,
                    'rows': len(df),
                    'columns': len(df.columns),
                    'file_size_kb': round(file_size / 1024, 2)
                }
            )
        
        except Exception as e:
            return self.handle_error('export_to_json', e)
    
    def generate_pdf_report(
        self,
        df: pd.DataFrame,
        file_path: str,
        title: str = "Data Analysis Report",
        summary_stats: Optional[Dict] = None
    ) -> Dict[str, Any]:
        """
        Generate a PDF summary report.
        
        Args:
            df: DataFrame to include in report
            file_path: Output file path
            title: Report title
            summary_stats: Optional summary statistics to include
            
        Returns:
            Result dictionary
        """
        try:
            is_valid, error = self.validate_dataframe(df)
            if not is_valid:
                return self.create_result(False, message=error)
            
            # Ensure directory exists
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            
            # Create PDF
            doc = SimpleDocTemplate(file_path, pagesize=letter)
            elements = []
            styles = getSampleStyleSheet()
            
            # Title
            elements.append(Paragraph(title, styles['Title']))
            elements.append(Spacer(1, 12))
            
            # Summary statistics
            if summary_stats:
                elements.append(Paragraph("Summary Statistics", styles['Heading2']))
                elements.append(Spacer(1, 6))
                
                for key, value in summary_stats.items():
                    elements.append(Paragraph(f"<b>{key}:</b> {value}", styles['Normal']))
                
                elements.append(Spacer(1, 12))
            
            # Data preview (first 20 rows)
            elements.append(Paragraph("Data Preview (First 20 Rows)", styles['Heading2']))
            elements.append(Spacer(1, 6))
            
            preview_df = df.head(20)
            
            # Create table data
            table_data = [preview_df.columns.tolist()]
            for _, row in preview_df.iterrows():
                table_data.append([str(val)[:30] for val in row.values])  # Truncate long values
            
            # Create table
            table = Table(table_data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            elements.append(table)
            
            # Build PDF
            doc.build(elements)
            
            file_size = os.path.getsize(file_path)
            
            self.log_operation('generate_pdf_report', {
                'file_path': file_path,
                'rows': len(df),
                'file_size_kb': file_size / 1024
            })
            
            return self.create_result(
                success=True,
                message=f"Generated PDF report: {file_path}",
                metadata={
                    'file_path': file_path,
                    'rows': len(df),
                    'file_size_kb': round(file_size / 1024, 2)
                }
            )
        
        except Exception as e:
            return self.handle_error('generate_pdf_report', e)
    
    def generate_word_report(
        self,
        df: pd.DataFrame,
        file_path: str,
        title: str = "Data Analysis Report",
        summary_stats: Optional[Dict] = None
    ) -> Dict[str, Any]:
        """
        Generate a Word document report.
        
        Args:
            df: DataFrame to include in report
            file_path: Output file path
            title: Report title
            summary_stats: Optional summary statistics to include
            
        Returns:
            Result dictionary
        """
        try:
            is_valid, error = self.validate_dataframe(df)
            if not is_valid:
                return self.create_result(False, message=error)
            
            # Ensure directory exists
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            
            # Create Word document
            doc = Document()
            
            # Title
            doc.add_heading(title, 0)
            
            # Summary statistics
            if summary_stats:
                doc.add_heading('Summary Statistics', level=1)
                for key, value in summary_stats.items():
                    doc.add_paragraph(f"{key}: {value}")
            
            # Data preview
            doc.add_heading('Data Preview (First 20 Rows)', level=1)
            
            preview_df = df.head(20)
            
            # Create table
            table = doc.add_table(rows=1, cols=len(preview_df.columns))
            table.style = 'Light Grid Accent 1'
            
            # Header row
            header_cells = table.rows[0].cells
            for i, column in enumerate(preview_df.columns):
                header_cells[i].text = str(column)
            
            # Data rows
            for _, row in preview_df.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row.values):
                    row_cells[i].text = str(value)[:50]  # Truncate long values
            
            # Save document
            doc.save(file_path)
            
            file_size = os.path.getsize(file_path)
            
            self.log_operation('generate_word_report', {
                'file_path': file_path,
                'rows': len(df),
                'file_size_kb': file_size / 1024
            })
            
            return self.create_result(
                success=True,
                message=f"Generated Word report: {file_path}",
                metadata={
                    'file_path': file_path,
                    'rows': len(df),
                    'file_size_kb': round(file_size / 1024, 2)
                }
            )
        
        except Exception as e:
            return self.handle_error('generate_word_report', e)
